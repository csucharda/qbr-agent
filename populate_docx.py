import json
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load structured JSON
with open("QBR_Q1_FY26_Lending_Approval_to_Funding.json", "r") as f:
    data = json.load(f)

# Load template
doc = Document("templates/202604_Domain-Adapted_QBRmemo_Template vSHARE.docx")

def set_cell(cell, text, *, white=False, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)

def quarter_num(q):
    return int(str(q).upper().replace("Q", ""))

def next_quarter_label(q):
    n = quarter_num(q)
    return f"Q{1 if n == 4 else n + 1}"

def parse_number(value):
    if value is None:
        return None
    s = str(value).strip()
    if s in ("", "-", "N/A", "To be confirmed"):
        return None
    s = s.replace("%", "").replace("pts", "").replace("pt", "")
    s = s.replace("+", "").replace("~", "").replace("cumulative", "")
    try:
        return float(s)
    except ValueError:
        return None

def set_dot(cell, actual, target):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("●")
    a = parse_number(actual)
    t = parse_number(target)
    if a is not None and t is not None and a >= t:
        run.font.color.rgb = RGBColor(0, 128, 0)
    else:
        run.font.color.rgb = RGBColor(192, 0, 0)

def set_left_cell(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("" if text is None else str(text))
    return run

def set_cell_text(cell, text):
    # Clear existing cell content
    cell.text = ""
    paragraphs = text.split("\n")
    cell.paragraphs[0].text = paragraphs[0] if paragraphs else ""
    for line in paragraphs[1:]:
        cell.add_paragraph(line)

def replace_in_paragraph(paragraph, old, new):
    if old in paragraph.text:
        paragraph.text = paragraph.text.replace(old, new)

def replace_everywhere(doc, old, new):
    # Normal paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p, old, new)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, old, new)

    # Headers and footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p, old, new)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p, old, new)

# Header fields
h = data["header"]

quarter_label = f"{h['quarter']} {h['fiscal_year']}"  # Q3 FY26

replace_everywhere(doc, "Qx FYxx", quarter_label)
replace_everywhere(doc, "FYxx", h["fiscal_year"])
replace_everywhere(doc, "[Domain]", h["domain"])
replace_everywhere(doc, "[Month] xx, xxxx", h["date"])
replace_everywhere(doc, "xxxxx xxxxx", h["domain_lead"])
replace_everywhere(doc, "Insert Domain Mission", h["mission"])
replace_everywhere(doc, "[Domain name]: Insert Domain Mission", h["mission"])

# Executive summary
es = data["executive_summary"]

learning_text = "\n".join([f"- {lp}" for lp in es["learning_points"]])
opportunity_text = "\n".join([f"- {op}" for op in es["next_opportunities"]])

summary_intro = f"The key priority for last quarter was to focus on {es['priority_last_quarter']}."

# Find executive summary table rows by their left-hand labels
for table in doc.tables:
    for row in table.rows:
        label = row.cells[0].text.strip()

        if label.startswith("The key priority for last quarter"):
            set_cell_text(row.cells[1], summary_intro)

        elif label.startswith("Key learnings this quarter were"):
            set_cell_text(row.cells[1], learning_text)

        elif label.startswith("Our next opportunities are"):
            set_cell_text(row.cells[1], opportunity_text)

# OKRs
okrs = data["okrs"]

okr_table = doc.tables[3]
row_idx = 1  # skip header row

for obj in okrs:
    objective = obj["objective"]
    upstream = obj.get("upstream_linkage", "")

    for kr in obj["key_results"]:
        kr_text = kr["kr"]
        status = kr.get("status", "")
        support = kr.get("support_needed", "")

        if row_idx >= len(okr_table.rows):
            break

        row = okr_table.rows[row_idx]

        row.cells[0].text = upstream
        row.cells[1].text = objective
        row.cells[2].text = kr_text
        row.cells[3].text = status
        row.cells[4].text = support

        row_idx += 1

# KPI dashboard
kpi_table = doc.tables[4]

current_q = h["quarter"]
next_q = next_quarter_label(h["quarter"])
fy = h["fiscal_year"]

kpi_header_row = kpi_table.rows[0]

# Force KPI header — correct indices based on merged cell structure
header_row = kpi_table.rows[0]

# Qx merged cell (cells[2,3,4] share same underlying element)
set_cell(header_row.cells[2], current_q, white=True, bold=True)

# Qx+1 (next column, not part of the Qx merge)
set_cell(header_row.cells[5], next_q, white=True, bold=True)

# FYxx merged cell (cells[6,7] share same underlying element — set both explicitly)
set_cell(header_row.cells[6], fy, white=True, bold=True)
set_cell(header_row.cells[7], fy, white=True, bold=True)

set_cell(kpi_table.rows[0].cells[8], "Upstream KPI Impact", white=True, bold=True)

# Second header row
set_cell(kpi_table.rows[1].cells[2], "Actual", bold=True)
set_cell(kpi_table.rows[1].cells[4], "Target", bold=True)
set_cell(kpi_table.rows[1].cells[5], "Target", bold=True)
set_cell(kpi_table.rows[1].cells[6], "Actuals\nto Date", bold=True)
set_cell(kpi_table.rows[1].cells[7], "Year End", bold=True)

# Replace any remaining placeholder text in header cells by value
for row in kpi_table.rows:
    for cell in row.cells:
        text = cell.text.strip()

        if text == "Qx":
            set_cell(cell, current_q, white=True, bold=True)

        elif text == "Qx+1":
            set_cell(cell, next_q, white=True, bold=True)

        elif text == "FYxx":
            set_cell(cell, fy, white=True, bold=True)

# Body rows
start_row = 2
for i, kpi in enumerate(data["kpis"]):
    row_idx = start_row + i
    if row_idx >= len(kpi_table.rows):
        break

    row = kpi_table.rows[row_idx]

    category = kpi.get("category", "")
    name = kpi.get("name", "")
    actual = kpi.get("q_current") or kpi.get("actuals_to_date") or ""
    q_next_target = kpi.get("q_next_target") or ""
    fy_target = kpi.get("fy_target") or ""
    actuals_to_date = kpi.get("actuals_to_date") or actual
    year_end = kpi.get("year_end") or fy_target
    upstream = kpi.get("upstream_kpi_impact") or ""

    set_cell(row.cells[0], category)
    set_cell(row.cells[1], name)
    set_cell(row.cells[2], actual)
    set_dot(row.cells[3], actual, q_next_target)
    set_cell(row.cells[4], q_next_target)
    set_cell(row.cells[5], fy_target)
    set_cell(row.cells[6], actuals_to_date)
    set_cell(row.cells[7], year_end)
    set_cell(row.cells[8], upstream)

# Section 1.4 Key Issues & Dependencies
dep_table = doc.tables[5]
deps = data.get("dependencies", [])

start_row = 1  # row 0 is the header
for i, dep in enumerate(deps):
    row_idx = start_row + i
    if row_idx >= len(dep_table.rows):
        break

    row = dep_table.rows[row_idx]

    set_left_cell(row.cells[0], dep.get("tracking_id", "TBC"))
    set_left_cell(row.cells[1], dep.get("constraint", ""))
    set_left_cell(row.cells[2], dep.get("required_action", ""))
    set_left_cell(row.cells[3], dep.get("decision_needed_by", ""))

# Section 1.5 Risks
risk_table = doc.tables[6]
risks = data.get("risks", [])

start_row = 1  # skip header
for i, risk in enumerate(risks):
    row_idx = start_row + i
    if row_idx >= len(risk_table.rows):
        break

    row = risk_table.rows[row_idx]

    row.cells[0].text = risk.get("risk", "")
    row.cells[1].text = risk.get("ask", "")

# Remove empty or placeholder rows from risks table
for i in range(len(risk_table.rows) - 1, 0, -1):
    row = risk_table.rows[i]
    row_text = " ".join(cell.text.strip() for cell in row.cells)

    if row_text == "" or "xxxx" in row_text.lower():
        risk_table._element.remove(row._element)

# Remove empty or placeholder rows from dependencies table
for i in range(len(dep_table.rows) - 1, 0, -1):  # backwards, skip header
    row = dep_table.rows[i]
    row_text = " ".join(cell.text.strip() for cell in row.cells)

    if row_text == "" or "xxxx" in row_text.lower():
        dep_table._element.remove(row._element)

# Appendix header
replace_everywhere(doc, "Appendix: Qx Initiatives", f"Appendix: {h['quarter']} Initiatives")

# Appendix: Initiatives
init_table = doc.tables[7]
initiatives = data.get("initiatives", [])

start_row = 1  # skip header row
for i, item in enumerate(initiatives):
    row_idx = start_row + i
    if row_idx >= len(init_table.rows):
        break

    row = init_table.rows[row_idx]

    jira = item.get("jira_planview_id", "")
    init_id = item.get("initiative_id", "")
    row.cells[0].text = f"{init_id}\n{jira}" if jira else str(init_id)
    row.cells[1].text = str(item.get("crew", ""))
    row.cells[2].text = str(item.get("initiative", ""))
    row.cells[3].text = str(item.get("outcome", ""))
    row.cells[4].text = str(item.get("relevant_okrs", ""))
    row.cells[5].text = str(item.get("confidence", ""))
    row.cells[6].text = str(item.get("key_risks_dependencies", ""))

# Remove unused placeholder rows from initiatives table
for i in range(len(init_table.rows) - 1, 0, -1):
    row = init_table.rows[i]
    row_text = " ".join(cell.text.strip() for cell in row.cells)
    if row_text == "" or "xxxx" in row_text.lower():
        init_table._element.remove(row._element)

# Save output
doc.save("QBR_Output.docx")
print("Word document generated: QBR_Output.docx")
